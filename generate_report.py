"""
课程评价系统 - 项目报告生成脚本
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

doc = Document()

# 页面设置
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)

# ============================================================
# 封面
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('人工智能通识课·实践项目')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('课程评价系统')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 102, 153)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_text = '''项目类别：第一类——智能体编程应用开发
技术栈：Python Flask + SQLite + HTML/CSS/JS + ECharts
GitHub：https://github.com/cty23/project'''
run = info.add_run(info_text)
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author.add_run('2026年5月')
run.font.size = Pt(14)

doc.add_page_break()

# ============================================================
# 一、任务阐述
# ============================================================
doc.add_heading('一、任务阐述', level=1)

doc.add_paragraph(
    '本项目来源于《人工智能通识课》实践项目任务书中"第一类：智能体编程应用开发"的要求。'
    '任务是利用AI编程助手协同开发一个真实可用的Web应用程序，系统需基于角色进行权限划分，并实现核心业务逻辑。'
)

doc.add_paragraph('核心硬性指标包括：')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('技术栈：前端（HTML/JS），后端（Python Flask），数据库（SQLite）')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('角色要求：实现至少2个角色（学生、教师、管理员）')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('功能要求：每个角色拥有至少3个核心功能')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('代码规范：完整代码、README.md运行说明及requirements.txt依赖配置')

doc.add_paragraph(
    '经过选题比较，本项目选择"课程评价系统"作为开发主题。'
    '该系统面向高校教学场景，允许学生对已选课程进行打分和文字评价，'
    '教师可以查看授课评价的统计数据和详细反馈，管理员负责课程和用户管理并查看全校数据仪表板。'
    '系统实现了完整的权限分流，不同角色登录后呈现完全不同的功能界面，'
    '并集成了ECharts图表库用于数据可视化展示。'
)

# ============================================================
# 二、背景调研
# ============================================================
doc.add_heading('二、背景调研', level=1)

doc.add_heading('2.1 课程评价的现实意义', level=2)
doc.add_paragraph(
    '课程评价是高等教育教学质量保障体系的重要组成部分。'
    '传统纸质评教方式存在效率低、统计难、反馈滞后等问题。'
    '构建一套线上课程评价系统，可以实现评价数据的实时采集、自动统计和可视化呈现，'
    '帮助教师及时了解教学效果、调整教学策略，为教学管理部门提供数据支撑。'
)

doc.add_heading('2.2 技术选型依据', level=2)
doc.add_paragraph(
    '后端选择Python Flask框架，原因如下：Flask轻量灵活，学习曲线平缓，'
    '适合中小型Web应用的快速开发；Python生态丰富，便于集成数据处理和导出功能。'
    '数据库选择SQLite，零配置、无需独立服务进程，适合教学演示和轻量部署场景。'
    '前端采用纯HTML/CSS/JS方案，避免引入前端框架的额外学习成本；'
    '数据可视化采用ECharts，是国内高校项目中广泛使用的图表库，文档完善、示例丰富。'
)

doc.add_heading('2.3 相关系统分析', level=2)
doc.add_paragraph(
    '目前市面上主流的教务管理系统（如正方、青果等）虽然包含评教模块，'
    '但多为大型商业系统，功能冗余、定制困难。'
    '部分高校使用问卷星等通用表单工具进行评教，但是数据管理分散、分析能力有限。'
    '本项目旨在开发一个轻量、专注、易部署的课程评价系统，作为现有方案的有效补充。'
)

# ============================================================
# 三、项目设计文档
# ============================================================
doc.add_heading('三、项目设计文档', level=1)

doc.add_heading('3.1 系统架构', level=2)
doc.add_paragraph(
    '系统采用经典的B/S（Browser/Server）三层架构：'
)

p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('表现层（Frontend）：纯HTML/CSS/JS页面，包含登录页及学生、教师、管理员三个功能界面')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('业务逻辑层（Backend）：Python Flask框架，处理路由、认证、业务逻辑和API接口')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('数据层（Database）：SQLite数据库，存储用户、课程、选课和评价数据')

doc.add_paragraph('系统架构图如下（截图位置）：')
doc.add_paragraph('[此处插入系统架构图]')

doc.add_heading('3.2 数据库设计', level=2)
doc.add_paragraph('数据库包含4张核心表：')

# users表
doc.add_heading('users（用户表）', level=3)
table = doc.add_table(rows=8, cols=4, style='Light Grid Accent 1')
headers = ['字段', '类型', '说明', '备注']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data = [
    ('id', 'INTEGER', '主键', '自增'),
    ('username', 'TEXT', '用户名', '唯一，登录用'),
    ('password_hash', 'TEXT', '密码哈希', 'werkzeug加密'),
    ('role', 'TEXT', '角色', 'student/teacher/admin'),
    ('display_name', 'TEXT', '显示名称', ''),
    ('email', 'TEXT', '邮箱', ''),
    ('created_at', 'TIMESTAMP', '创建时间', '默认当前时间'),
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i + 1].cells[j].text = val

doc.add_paragraph()

# courses表
doc.add_heading('courses（课程表）', level=3)
table = doc.add_table(rows=5, cols=4, style='Light Grid Accent 1')
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data = [
    ('id', 'INTEGER', '主键', '自增'),
    ('name', 'TEXT', '课程名称', ''),
    ('description', 'TEXT', '课程描述', ''),
    ('teacher_id', 'INTEGER', '授课教师ID', '外键→users.id'),
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i + 1].cells[j].text = val

doc.add_paragraph()

# enrollments表
doc.add_heading('enrollments（选课表）', level=3)
table = doc.add_table(rows=4, cols=4, style='Light Grid Accent 1')
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data = [
    ('id', 'INTEGER', '主键', '自增'),
    ('student_id', 'INTEGER', '学生ID', '外键→users.id'),
    ('course_id', 'INTEGER', '课程ID', '外键→courses.id'),
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i + 1].cells[j].text = val

doc.add_paragraph()

# evaluations表
doc.add_heading('evaluations（评价表）', level=3)
table = doc.add_table(rows=6, cols=4, style='Light Grid Accent 1')
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data = [
    ('id', 'INTEGER', '主键', '自增'),
    ('student_id', 'INTEGER', '学生ID', '外键→users.id'),
    ('course_id', 'INTEGER', '课程ID', '外键→courses.id'),
    ('rating', 'INTEGER', '评分', '1-5星'),
    ('comment', 'TEXT', '文字评语', ''),
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i + 1].cells[j].text = val

doc.add_paragraph()

doc.add_heading('3.3 功能设计', level=2)

# 角色功能表
doc.add_paragraph('三个角色的核心功能如下：')
table = doc.add_table(rows=4, cols=4, style='Light Grid Accent 1')
headers = ['角色', '功能1', '功能2', '功能3']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

data = [
    ('学生', '浏览课程列表并选课', '提交课程评价（1-5星+文字）', '查看个人评价历史'),
    ('教师', '查看授课评价统计（均分/分布）', '浏览学生详细评价列表', '导出评价数据为CSV'),
    ('管理员', '课程管理（增删改）', '用户管理（查看/修改角色）', '全校数据仪表板（ECharts）'),
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i + 1].cells[j].text = val

doc.add_heading('3.4 权限控制设计', level=2)
doc.add_paragraph(
    '系统基于角色的访问控制（RBAC）模型实现权限管理。'
    '登录成功后，后端根据用户角色返回对应token，重定向到角色专属页面。'
    '关键设计点：'
)
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('每个HTML页面在加载时通过fetch调用/api/me接口验证身份和角色')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('后端使用装饰器@login_required和@require_role进行接口级别的权限拦截')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('API层对每次请求重新校验角色，防止越权访问')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('前端根据角色动态渲染菜单和操作按钮，实现界面级权限分流')

doc.add_heading('3.5 项目目录结构', level=2)
doc.add_paragraph('项目采用前后端分离的目录组织方式：')

code = '''course-evaluation-system/
├── backend/
│   ├── app.py          # Flask主程序（路由、API）
│   ├── models.py       # 数据库模型与初始化
│   └── auth.py         # 登录验证与权限装饰器
├── frontend/
│   ├── index.html      # 登录页
│   ├── student.html    # 学生功能界面
│   ├── teacher.html    # 教师功能界面
│   └── admin.html      # 管理员功能界面
├── database/
│   └── init.sql        # 数据库初始化脚本
├── requirements.txt    # Python依赖
└── README.md           # 运行说明'''
p = doc.add_paragraph()
run = p.add_run(code)
run.font.name = 'Consolas'
run.font.size = Pt(9)

# ============================================================
# 四、结果展示
# ============================================================
doc.add_heading('四、结果展示', level=1)

doc.add_heading('4.1 登录页面', level=2)
doc.add_paragraph(
    '登录页为系统入口，采用居中卡片式布局。用户输入用户名和密码后，'
    '系统自动识别角色并跳转至对应功能页面。'
    '预置测试账号包括管理员（admin）、3位教师（teacher1-teacher3）、5位学生（student1-student5）。'
)
doc.add_paragraph('[此处插入登录页截图]')

doc.add_heading('4.2 学生端功能', level=2)
doc.add_paragraph('学生登录后可执行以下操作：')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('浏览全部课程，点击"选课"按钮加入课程')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('对已选课程进行1-5星打分并填写文字评语')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('在"我的评价"区域查看历史评价记录及课程信息')
doc.add_paragraph('[此处插入学生端截图]')

doc.add_heading('4.3 教师端功能', level=2)
doc.add_paragraph('教师登录后可执行以下操作：')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('查看所授课程的平均评分和评分分布柱状图')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('浏览学生对课程的详细评价列表')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('点击导出按钮将评价数据下载为CSV文件')
doc.add_paragraph('[此处插入教师端截图]')

doc.add_heading('4.4 管理员端功能', level=2)
doc.add_paragraph('管理员登录后可执行以下操作：')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('课程管理：添加、编辑、删除课程，分配授课教师')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('用户管理：查看全部用户列表，修改用户角色')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('数据仪表板：课程平均分排名柱状图、评分分布饼图、选课统计等')
doc.add_paragraph('[此处插入管理员端截图]')

doc.add_heading('4.5 数据可视化', level=2)
doc.add_paragraph(
    '系统集成ECharts库实现数据可视化，包括：课程平均评分排名柱状图、'
    '各分数段评价数量饼图、课程选课人数横向对比图。'
    '所有图表数据通过后端API实时查询数据库生成，支持交互式tooltip悬停查看详情。'
)
doc.add_paragraph('[此处插入数据可视化截图]')

# ============================================================
# 五、GitHub仓库证明
# ============================================================
doc.add_heading('五、GitHub仓库证明', level=1)

doc.add_paragraph('项目代码已托管至GitHub公开仓库：')
p = doc.add_paragraph()
run = p.add_run('仓库地址：https://github.com/cty23/project')
run.bold = True

doc.add_paragraph('仓库包含以下内容：')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('backend/：完整后端代码（Flask路由、数据库模型、认证模块）')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('frontend/：4个前端HTML页面（登录、学生、教师、管理员）')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('database/：SQLite数据库初始化脚本')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('requirements.txt：Python依赖清单')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('README.md：详细的安装、配置和运行说明')

doc.add_paragraph('[此处插入GitHub仓库首页截图]')

# ============================================================
# 六、展望与总结
# ============================================================
doc.add_heading('六、展望与总结', level=1)

doc.add_heading('6.1 项目总结', level=2)
doc.add_paragraph(
    '本项目完成了课程评价系统的全部核心功能开发，实现了学生选课评教、'
    '教师查看统计分析、管理员系统管理的完整业务流程。'
    '系统达到了任务书中全部硬性指标要求：使用了指定的技术栈（Flask+SQLite+HTML/JS），'
    '实现了3个角色的权限控制，每个角色至少3项核心功能，'
    '包含完整的README运行说明和依赖配置文件。'
)

doc.add_paragraph('在加分项方面，本项目实现了：')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('完整的权限控制——三个角色登录后显示完全不同的功能界面')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('数据可视化——集成了ECharts图表库，在管理员仪表板中展示多项统计图表')

doc.add_heading('6.2 未来展望', level=2)
doc.add_paragraph('未来可从以下方向进行迭代升级：')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('部署到云服务器（如阿里云、腾讯云），实现公网访问')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('增加匿名评价功能，保护学生隐私')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('引入评价模型分析（情感分析/NLP）自动识别关键反馈')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('对接学校统一身份认证系统（SSO/CAS）')
p = doc.add_paragraph()
p.style = 'List Bullet'
p.add_run('增加评价提醒和自动催办功能')

doc.add_heading('6.3 开发体会', level=2)
doc.add_paragraph(
    '通过本项目，我深入实践了AI辅助编程的工作模式。'
    '在与AI编程助手的协作过程中，学习了如何清晰地描述需求、'
    '分解复杂任务、验证生成代码的正确性。'
    '项目开发从需求分析到代码实现再到部署交付，完整覆盖了软件工程的各个阶段，'
    '加深了对Web应用开发全流程的理解。'
    '同时也认识到，AI助手虽然能大幅提升开发效率，'
    '但开发者的系统设计能力、问题定位能力和质量标准把控能力仍然是不可替代的核心素养。'
)

# ============================================================
# 保存
# ============================================================
output_path = r'C:\Users\cty27\.openclaw\workspace\course-evaluation-system\课程评价系统-项目报告.docx'
doc.save(output_path)
print(f'报告已生成：{output_path}')
